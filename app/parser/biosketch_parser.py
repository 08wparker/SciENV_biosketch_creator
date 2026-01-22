"""Parser for NIH Biosketch Word documents."""

from __future__ import annotations
import re
from pathlib import Path
from typing import List, Union
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .models import (
    BiosketchData,
    Education,
    Position,
    Honor,
    Citation,
    Contribution,
    PersonalStatement,
    Grant
)
from .citation_parser import CitationParser


class BiosketchParser:
    """Parser for NIH Biosketch Word documents."""

    # Section header patterns
    SECTION_A_PATTERN = re.compile(r'^A\.\s*Personal\s*Statement', re.IGNORECASE)
    SECTION_B_PATTERN = re.compile(r'^B\.\s*Positions', re.IGNORECASE)
    SECTION_C_PATTERN = re.compile(r'^C\.\s*Contributions?\s*to\s*Science', re.IGNORECASE)

    # Header field patterns
    NAME_PATTERN = re.compile(r'^NAME:\s*(.+)$', re.IGNORECASE)
    ERA_COMMONS_PATTERN = re.compile(r'^eRA\s*COMMONS\s*USER\s*NAME[^:]*:\s*(.+)$', re.IGNORECASE)
    POSITION_TITLE_PATTERN = re.compile(r'^POSITION\s*TITLE:\s*(.+)$', re.IGNORECASE)

    def __init__(self, file_path: Union[str, Path]):
        self.file_path = Path(file_path)
        self.doc: DocxDocument = Document(str(self.file_path))
        self.data = BiosketchData()

    @staticmethod
    def iter_block_items(parent):
        """Iterate through paragraphs and tables in document order.

        This is crucial for maintaining the correct order of elements
        when parsing the document.
        """
        parent_elm = parent.element.body
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def parse(self) -> BiosketchData:
        """Parse the entire biosketch document."""
        self._parse_header()
        self._parse_education_table()
        self._parse_sections()
        return self.data

    def _parse_header(self):
        """Parse the header fields (Name, eRA Commons, Position Title)."""
        for para in self.doc.paragraphs[:15]:
            text = para.text.strip()
            if not text:
                continue

            # Try to match each header field
            name_match = self.NAME_PATTERN.match(text)
            if name_match:
                self.data.name = name_match.group(1).strip()
                continue

            era_match = self.ERA_COMMONS_PATTERN.match(text)
            if era_match:
                self.data.era_commons_username = era_match.group(1).strip()
                continue

            position_match = self.POSITION_TITLE_PATTERN.match(text)
            if position_match:
                self.data.position_title = position_match.group(1).strip()
                continue

    def _parse_education_table(self):
        """Parse the education/training table (first table in document)."""
        if not self.doc.tables:
            return

        table = self.doc.tables[0]

        # Skip header row
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 4 and cells[0]:  # Must have at least institution
                education = Education(
                    institution=cells[0],
                    degree=cells[1] if len(cells) > 1 else '',
                    completion_date=cells[2] if len(cells) > 2 else '',
                    field_of_study=cells[3] if len(cells) > 3 else ''
                )
                self.data.education.append(education)

    def _parse_sections(self):
        """Parse sections A, B, and C of the biosketch."""
        current_section = None
        section_content = []

        for para in self.doc.paragraphs:
            text = para.text.strip()

            # Check for section headers
            if self.SECTION_A_PATTERN.match(text):
                if current_section:
                    self._process_section(current_section, section_content)
                current_section = 'A'
                section_content = []
                continue
            elif self.SECTION_B_PATTERN.match(text):
                if current_section:
                    self._process_section(current_section, section_content)
                current_section = 'B'
                section_content = []
                continue
            elif self.SECTION_C_PATTERN.match(text):
                if current_section:
                    self._process_section(current_section, section_content)
                current_section = 'C'
                section_content = []
                continue

            # Collect content for current section
            if current_section and text:
                section_content.append(text)

        # Process the last section
        if current_section:
            self._process_section(current_section, section_content)

    def _process_section(self, section: str, content: List[str]):
        """Process a section's content based on section type."""
        if section == 'A':
            self._parse_section_a(content)
        elif section == 'B':
            self._parse_section_b(content)
        elif section == 'C':
            self._parse_section_c(content)

    def _parse_section_a(self, content: List[str]):
        """Parse Section A - Personal Statement."""
        personal_text = []
        grant_lines = []
        citations = []

        in_research_support = False
        in_citations = False

        for line in content:
            # Check for research support header
            if re.match(r'^Current\s*(and\s*recently\s*completed)?\s*research\s*support', line, re.IGNORECASE):
                in_research_support = True
                in_citations = False
                continue

            # Check for citations header
            if re.match(r'^Citations?:?\s*$', line, re.IGNORECASE):
                in_research_support = False
                in_citations = True
                continue

            if in_citations:
                # Parse as citation
                citation = CitationParser.parse_citation(line)
                citations.append(citation)
            elif in_research_support:
                # Collect grant lines for later parsing
                grant_lines.append(line)
            else:
                # Personal statement text
                personal_text.append(line)

        # Parse grants from collected lines
        grants = self._parse_grants(grant_lines)

        self.data.personal_statement = PersonalStatement(
            text=' '.join(personal_text),
            grants=grants,
            citations=citations
        )

    def _parse_grants(self, lines: List[str]) -> List[Grant]:
        """Parse research support/grant entries.

        Grants typically appear as:
        - Line 1: [Funder] [Number] [PI] (PI/role) [dates]
        - Line 2: Grant title
        - Optional Line 3: Role: Co-investigator (if not PI)

        Examples:
        NIH K08 HL150291    Parker (PI)    02/01/2020 - 01/31/2025
        Mending a Broken Heart Allocation System with Machine Learning

        NIH R01HL173037    Mayampurath (PI)    5/2024 – 03/2029
        Clinical Decision Support for Early Detection...
        Role: Co-investigator
        """
        grants = []
        current_grant = None

        # Pattern for grant header line: funder + number + PI + dates
        # e.g., "NIH K08 HL150291    Parker (PI)    02/01/2020 - 01/31/2025"
        grant_header_pattern = re.compile(
            r'^(NIH|NSF|DoD|VA|CDC|Greenwall|AHRQ|PCORI|[A-Z][a-z]+\s+Foundation)'  # Funder
            r'[:\s]+'
            r'([A-Z0-9\-\s]+?)'  # Grant number (or "No number")
            r'\s+'
            r'([A-Za-z,\s]+?)'  # PI name(s)
            r'\s*\(([^)]+)\)'  # Role in parentheses (PI, contact PI, etc.)
            r'\s*'
            r'(\d{1,2}/\d{1,4}\s*[-–]\s*\d{1,2}/\d{1,4}|\d{4}\s*[-–]\s*\d{4})?',  # Date range (optional)
            re.IGNORECASE
        )

        # Simpler pattern for lines that look like grant headers
        simple_grant_pattern = re.compile(
            r'^(NIH|NSF|DoD|VA|CDC|Greenwall|AHRQ|PCORI)',
            re.IGNORECASE
        )

        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # Check if this is a Role: line
            role_match = re.match(r'^Role:\s*(.+)$', line, re.IGNORECASE)
            if role_match and current_grant:
                current_grant.role = role_match.group(1).strip()
                i += 1
                continue

            # Try to match grant header pattern
            header_match = grant_header_pattern.match(line)
            if header_match:
                # Save previous grant if exists
                if current_grant:
                    grants.append(current_grant)

                current_grant = Grant(
                    funder=header_match.group(1).strip(),
                    number=header_match.group(2).strip() if header_match.group(2) else "",
                    pi=header_match.group(3).strip() if header_match.group(3) else "",
                    role=header_match.group(4).strip() if header_match.group(4) else "PI",
                    dates=header_match.group(5).strip() if header_match.group(5) else ""
                )
                i += 1
                continue

            # Check if this looks like a grant header (starts with known funder)
            if simple_grant_pattern.match(line):
                # Save previous grant
                if current_grant:
                    grants.append(current_grant)

                # Try to parse the line manually
                current_grant = self._parse_grant_line(line)
                i += 1
                continue

            # Otherwise, this might be a grant title
            if current_grant and not current_grant.title:
                current_grant.title = line
            elif current_grant:
                # Append to title if it continues
                current_grant.title += " " + line

            i += 1

        # Don't forget the last grant
        if current_grant:
            grants.append(current_grant)

        return grants

    def _parse_grant_line(self, line: str) -> Grant:
        """Parse a single grant header line into components."""
        grant = Grant()

        # Try to extract dates (at the end typically)
        date_pattern = re.compile(r'(\d{1,2}/\d{1,4}\s*[-–]\s*\d{1,2}/\d{1,4})')
        date_match = date_pattern.search(line)
        if date_match:
            grant.dates = date_match.group(1).strip()
            line = line[:date_match.start()] + line[date_match.end():]

        # Try to extract PI info (in parentheses)
        pi_pattern = re.compile(r'([A-Za-z,\s]+)\s*\(([^)]+)\)')
        pi_match = pi_pattern.search(line)
        if pi_match:
            grant.pi = pi_match.group(1).strip()
            grant.role = pi_match.group(2).strip()
            line = line[:pi_match.start()] + line[pi_match.end():]

        # Split remaining into funder and number
        parts = line.strip().split(None, 1)
        if parts:
            grant.funder = parts[0].strip()
            if len(parts) > 1:
                grant.number = parts[1].strip()

        return grant

    def _parse_section_b(self, content: List[str]):
        """Parse Section B - Positions and Honors."""
        in_positions = True
        in_honors = False

        for line in content:
            # Check for honors header
            if re.match(r'^Honors?\s*$', line, re.IGNORECASE):
                in_positions = False
                in_honors = True
                continue

            # Skip section sub-headers
            if re.match(r'^Positions?\s*(and\s*Scientific\s*Appointments?)?\s*$', line, re.IGNORECASE):
                continue

            if in_honors:
                # Parse honor entry: "Year   Description"
                honor_match = re.match(r'^(\d{4})\s+(.+)$', line)
                if honor_match:
                    self.data.honors.append(Honor(
                        year=honor_match.group(1),
                        description=honor_match.group(2).strip()
                    ))
            elif in_positions:
                # Parse position entry: "Dates   Title, Institution" or "Dates   Title"
                # Common formats:
                # "2021-Present   Assistant Professor of Medicine, University of Chicago"
                # "2015-2019      Fellow, Pulmonary and Critical Care Medicine, University of Chicago"
                position_match = re.match(r'^(\d{4}[-–]\d{4}|\d{4}[-–]Present)\s+(.+)$', line, re.IGNORECASE)
                if position_match:
                    dates = position_match.group(1)
                    rest = position_match.group(2).strip()

                    # Try to split title and institution by comma
                    # The institution is typically the last comma-separated part
                    parts = rest.split(',')
                    if len(parts) >= 2:
                        # Last part is likely institution
                        institution = parts[-1].strip()
                        title = ','.join(parts[:-1]).strip()
                    else:
                        title = rest
                        institution = ''

                    self.data.positions.append(Position(
                        dates=dates,
                        title=title,
                        institution=institution
                    ))

    def _parse_section_c(self, content: List[str]):
        """Parse Section C - Contributions to Science."""
        contributions = []
        current_narrative = []
        current_citations = []

        # Pattern to detect contribution headers (numbered or titled)
        contribution_header = re.compile(r'^(\d+\.?\s+)?[A-Z]')

        for i, line in enumerate(content):
            # Skip the "Complete List of Published Work" line at the end
            if 'Complete List of Published Work' in line:
                continue

            # Check if this line looks like a citation
            is_citation = CitationParser.is_citation_line(line)

            # Check if this looks like a new contribution header
            # (starts with number or capital letter, is not a citation)
            is_new_contribution = (
                contribution_header.match(line) and
                not is_citation and
                len(line) > 20 and
                not line.startswith('Role:')
            )

            if is_new_contribution and (current_narrative or current_citations):
                # Save previous contribution
                if current_narrative:
                    contributions.append(Contribution(
                        narrative=' '.join(current_narrative),
                        citations=current_citations.copy()
                    ))
                current_narrative = [line]
                current_citations = []
            elif is_citation:
                citation = CitationParser.parse_citation(line)
                current_citations.append(citation)
            else:
                # Continue building narrative
                current_narrative.append(line)

        # Don't forget the last contribution
        if current_narrative:
            contributions.append(Contribution(
                narrative=' '.join(current_narrative),
                citations=current_citations.copy()
            ))

        self.data.contributions = contributions


def parse_biosketch(file_path: Union[str, Path]) -> BiosketchData:
    """Convenience function to parse a biosketch file."""
    parser = BiosketchParser(file_path)
    return parser.parse()
